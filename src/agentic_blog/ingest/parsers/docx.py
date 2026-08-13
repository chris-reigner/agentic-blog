"""DOCX parser: python-docx preferred, hardened stdlib zipfile fallback.

The zipfile fallback validates the embedded XML against DTD/ENTITY declarations
to defend against XXE and billion-laughs entity expansion.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from agentic_blog.contracts import RawDocument
from agentic_blog.ingest.parsers.base import ExtractionError, source_id_for

EXTENSIONS = {".docx"}

_UNSAFE_XML = re.compile(r"<!DOCTYPE|<!ENTITY", re.IGNORECASE)


def _validate_xml_safety(xml: str) -> None:
    """Reject DTD/ENTITY constructs used in XXE / billion-laughs attacks."""
    if _UNSAFE_XML.search(xml):
        raise ExtractionError(
            "DOCX contains disallowed DTD/ENTITY declarations; refusing to parse."
        )


def _with_python_docx(path: str) -> str | None:
    try:
        import docx
    except ImportError:
        return None
    try:
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)
    except Exception:
        return None


def _with_zipfile(path: str) -> str | None:
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    except (KeyError, zipfile.BadZipFile, OSError):
        return None
    _validate_xml_safety(xml)
    # Insert breaks for paragraph and line-break tags, then strip all markup.
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"<w:br\b[^>]*/>", "\n", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    return text


class DocxParser:
    def can_parse(self, origin: str) -> bool:
        return Path(origin).suffix.lower() in EXTENSIONS

    def parse(self, origin: str) -> RawDocument:
        for extractor, name in ((_with_python_docx, "python-docx"), (_with_zipfile, "zipfile")):
            text = extractor(origin)
            if text and text.strip():
                return RawDocument(
                    source_id=source_id_for(origin),
                    origin=origin,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    text=text,
                    title=Path(origin).stem,
                    metadata={"parser": name},
                )
        raise ExtractionError(f"Could not extract text from DOCX: {origin}")
